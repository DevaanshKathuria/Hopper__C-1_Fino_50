"""Build final report and presentation PDF artifacts for submission packaging.

The source of truth remains the Markdown files in reports/. This script creates
submission-friendly DOCX/PDF artifacts without inventing analytical claims.
"""

from __future__ import annotations

import re
from pathlib import Path
from textwrap import shorten

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


REPO_ROOT = Path(__file__).resolve().parents[1]
REPORT_MD = REPO_ROOT / "reports" / "project_report.md"
PRESENTATION_MD = REPO_ROOT / "reports" / "presentation_outline.md"
REPORT_DOCX = REPO_ROOT / "reports" / "project_report.docx"
REPORT_PDF = REPO_ROOT / "reports" / "project_report.pdf"
PRESENTATION_PDF = REPO_ROOT / "reports" / "presentation.pdf"

KEY_CHARTS = [
    "yearly_market_trend.png",
    "sector_return_comparison.png",
    "risk_return_segments.png",
    "covid_crash_recovery.png",
]


def strip_inline_markdown(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "")
    return text


def parse_markdown(path: Path) -> list[dict[str, object]]:
    blocks: list[dict[str, object]] = []
    table_buffer: list[str] = []
    bullet_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            blocks.append({"type": "table", "rows": table_buffer})
            table_buffer = []

    def flush_bullets() -> None:
        nonlocal bullet_buffer
        if bullet_buffer:
            blocks.append({"type": "bullets", "items": bullet_buffer})
            bullet_buffer = []

    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        if not line:
            flush_table()
            flush_bullets()
            continue

        if line.startswith("|") and line.endswith("|"):
            flush_bullets()
            if not re.match(r"^\|\s*-+", line):
                table_buffer.append(line)
            continue

        flush_table()
        if line.startswith("- "):
            bullet_buffer.append(strip_inline_markdown(line[2:]))
            continue

        flush_bullets()
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            blocks.append({"type": "heading", "level": level, "text": strip_inline_markdown(line.lstrip("#").strip())})
        elif re.match(r"^\d+\. ", line):
            blocks.append({"type": "paragraph", "text": strip_inline_markdown(line)})
        else:
            blocks.append({"type": "paragraph", "text": strip_inline_markdown(line)})

    flush_table()
    flush_bullets()
    return blocks


def table_rows(markdown_rows: list[str]) -> list[list[str]]:
    rows = []
    for row in markdown_rows:
        cells = [strip_inline_markdown(cell.strip()) for cell in row.strip("|").split("|")]
        rows.append(cells)
    return rows


def build_report_docx(blocks: list[dict[str, object]]) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(11)

    for block in blocks:
        kind = block["type"]
        if kind == "heading":
            text = str(block["text"])
            level = int(block["level"])
            if level == 1:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(text)
                run.bold = True
                run.font.size = Pt(20)
                document.add_paragraph()
            else:
                document.add_heading(text, level=min(level - 1, 3))
        elif kind == "paragraph":
            document.add_paragraph(str(block["text"]))
        elif kind == "bullets":
            for item in block["items"]:
                document.add_paragraph(str(item), style="List Bullet")
        elif kind == "table":
            rows = table_rows(block["rows"])
            if not rows:
                continue
            table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
            table.style = "Table Grid"
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    table.cell(i, j).text = cell
            document.add_paragraph()

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Selected Chart Appendix", level=1)
    for chart_name in KEY_CHARTS:
        chart_path = REPO_ROOT / "outputs" / "charts" / chart_name
        if chart_path.exists():
            document.add_paragraph(chart_name.replace("_", " ").replace(".png", "").title())
            document.add_picture(str(chart_path), width=Inches(6.8))

    document.save(REPORT_DOCX)


def report_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "Title",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=27,
            textColor=colors.HexColor("#102A43"),
            alignment=TA_CENTER,
            spaceAfter=18,
        ),
        "h1": ParagraphStyle(
            "Heading1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=18,
            textColor=colors.HexColor("#102A43"),
            spaceBefore=10,
            spaceAfter=6,
        ),
        "h2": ParagraphStyle(
            "Heading2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=14,
            textColor=colors.HexColor("#334E68"),
            spaceBefore=8,
            spaceAfter=4,
        ),
        "body": ParagraphStyle(
            "Body",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8.2,
            leading=10.8,
            alignment=TA_LEFT,
            spaceAfter=4,
        ),
        "small": ParagraphStyle(
            "Small",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=6.8,
            leading=8.2,
            spaceAfter=2,
        ),
    }


def pdf_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 7)
    canvas.setFillColor(colors.HexColor("#627D98"))
    canvas.drawString(0.55 * inch, 0.35 * inch, "Fino_50 (C-1) | NIFTY-50 Sectoral Performance & Risk Intelligence")
    canvas.drawRightString(A4[0] - 0.55 * inch, 0.35 * inch, f"Page {doc.page}")
    canvas.restoreState()


def build_report_pdf(blocks: list[dict[str, object]]) -> None:
    styles = report_styles()
    doc = SimpleDocTemplate(
        str(REPORT_PDF),
        pagesize=A4,
        leftMargin=0.5 * inch,
        rightMargin=0.5 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
    )
    story = []

    for block in blocks:
        kind = block["type"]
        if kind == "heading":
            text = str(block["text"])
            level = int(block["level"])
            if level == 1:
                story.append(Paragraph(text, styles["title"]))
            elif level == 2:
                story.append(Paragraph(text, styles["h1"]))
            else:
                story.append(Paragraph(text, styles["h2"]))
        elif kind == "paragraph":
            story.append(Paragraph(strip_inline_markdown(str(block["text"])), styles["body"]))
        elif kind == "bullets":
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(str(item), styles["body"]), leftIndent=10) for item in block["items"]],
                    bulletType="bullet",
                    leftIndent=14,
                )
            )
        elif kind == "table":
            rows = table_rows(block["rows"])
            if not rows:
                continue
            max_cols = max(len(row) for row in rows)
            padded = [row + [""] * (max_cols - len(row)) for row in rows]
            data = [[Paragraph(shorten(cell, width=80, placeholder="..."), styles["small"]) for cell in row] for row in padded]
            table = Table(data, repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9EAF7")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#102A43")),
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#BCCCDC")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 3),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 4))

    story.append(PageBreak())
    story.append(Paragraph("Selected Chart Appendix", styles["h1"]))
    for chart_index, chart_name in enumerate(KEY_CHARTS):
        chart_path = REPO_ROOT / "outputs" / "charts" / chart_name
        if chart_path.exists():
            story.append(
                KeepTogether(
                    [
                        Paragraph(chart_name.replace("_", " ").replace(".png", "").title(), styles["h2"]),
                        Image(str(chart_path), width=7.1 * inch, height=3.5 * inch),
                        Spacer(1, 8),
                    ]
                )
            )
            if chart_index != len(KEY_CHARTS) - 1:
                story.append(PageBreak())

    doc.build(story, onFirstPage=pdf_footer, onLaterPages=pdf_footer)


def presentation_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "cover": ParagraphStyle(
            "Cover",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=25,
            leading=31,
            textColor=colors.white,
            alignment=TA_LEFT,
        ),
        "slide_title": ParagraphStyle(
            "SlideTitle",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=26,
            textColor=colors.HexColor("#102A43"),
            spaceAfter=12,
        ),
        "slide_body": ParagraphStyle(
            "SlideBody",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=13,
            leading=17,
            textColor=colors.HexColor("#243B53"),
        ),
        "slide_small": ParagraphStyle(
            "SlideSmall",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=13,
            textColor=colors.HexColor("#52606D"),
        ),
    }


def parse_slide_outline() -> list[tuple[str, list[str]]]:
    slides: list[tuple[str, list[str]]] = []
    current_title = ""
    current_items: list[str] = []
    for line in PRESENTATION_MD.read_text(encoding="utf-8").splitlines():
        if line.startswith("## Slide "):
            if current_title:
                slides.append((current_title, current_items))
            current_title = strip_inline_markdown(line.replace("## ", ""))
            current_items = []
        elif line.startswith("- ") and current_title:
            current_items.append(strip_inline_markdown(line[2:]))
    if current_title:
        slides.append((current_title, current_items))
    return slides


def build_presentation_pdf() -> None:
    styles = presentation_styles()
    doc = SimpleDocTemplate(
        str(PRESENTATION_PDF),
        pagesize=landscape(A4),
        leftMargin=0.45 * inch,
        rightMargin=0.45 * inch,
        topMargin=0.45 * inch,
        bottomMargin=0.45 * inch,
    )
    story = []
    slides = parse_slide_outline()

    for index, (title, items) in enumerate(slides, start=1):
        if index == 1:
            cover_data = [
                [
                    Paragraph(
                        "NIFTY-50 Sectoral Performance & Risk Intelligence",
                        styles["cover"],
                    )
                ]
            ]
            cover = Table(cover_data, colWidths=[10.6 * inch], rowHeights=[5.1 * inch])
            cover.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#102A43")),
                        ("BOX", (0, 0), (-1, -1), 1, colors.HexColor("#102A43")),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 36),
                    ]
                )
            )
            story.append(cover)
            story.append(Spacer(1, 10))
            story.append(Paragraph("Fino_50 (C-1) | Finance / Stock Market Analytics", styles["slide_body"]))
        else:
            story.append(Paragraph(title, styles["slide_title"]))
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(item, styles["slide_body"]), leftIndent=12) for item in items],
                    bulletType="bullet",
                    leftIndent=20,
                )
            )
            if index in {6, 7, 8}:
                chart_name = {6: "yearly_market_trend.png", 7: "risk_return_segments.png", 8: "sector_return_comparison.png"}[index]
                chart_path = REPO_ROOT / "outputs" / "charts" / chart_name
                if chart_path.exists():
                    story.append(Spacer(1, 10))
                    story.append(Image(str(chart_path), width=6.6 * inch, height=3.1 * inch))
        if index != len(slides):
            story.append(PageBreak())

    doc.build(story)


def main() -> None:
    blocks = parse_markdown(REPORT_MD)
    build_report_docx(blocks)
    build_report_pdf(blocks)
    build_presentation_pdf()

    for path in [REPORT_DOCX, REPORT_PDF, PRESENTATION_PDF]:
        print(f"created: {path}")
    print(f"report_pdf_pages: {len(PdfReader(str(REPORT_PDF)).pages)}")
    print(f"presentation_pdf_pages: {len(PdfReader(str(PRESENTATION_PDF)).pages)}")


if __name__ == "__main__":
    main()
